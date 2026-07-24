"""Deterministic PDF and DOCX exporters for prose manuscripts."""

from __future__ import annotations

from datetime import UTC, datetime
from io import BytesIO
from typing import Any
from xml.sax.saxutils import escape
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document as create_document
from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY  # type: ignore[import-untyped]
from reportlab.lib.pagesizes import LETTER  # type: ignore[import-untyped]
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore[import-untyped]
from reportlab.lib.units import inch  # type: ignore[import-untyped]
from reportlab.pdfgen.canvas import Canvas  # type: ignore[import-untyped]
from reportlab.platypus import (  # type: ignore[import-untyped]
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

from open_hollywood_engine.rendering.contracts import ProseManuscript

_STATIC_TIMESTAMP = datetime(2000, 1, 1, tzinfo=UTC)
_ZIP_TIMESTAMP = (2000, 1, 1, 0, 0, 0)
_BLUE = RGBColor(0x2E, 0x74, 0xB5)
_DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)


def export_pdf(manuscript: ProseManuscript) -> bytes:
    """Export a stable, searchable US-Letter PDF."""
    output = BytesIO()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ManuscriptTitle",
        parent=styles["Title"],
        fontName="Times-Roman",
        fontSize=28,
        leading=34,
        alignment=TA_CENTER,
        spaceAfter=18,
    )
    byline_style = ParagraphStyle(
        "ManuscriptByline",
        parent=styles["Normal"],
        fontName="Times-Italic",
        fontSize=12,
        leading=16,
        alignment=TA_CENTER,
    )
    scene_style = ParagraphStyle(
        "SceneHeading",
        parent=styles["Heading2"],
        fontName="Times-Bold",
        fontSize=15,
        leading=19,
        spaceBefore=8,
        spaceAfter=16,
        keepWithNext=True,
    )
    body_style = ParagraphStyle(
        "ManuscriptBody",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=11,
        leading=16,
        alignment=TA_JUSTIFY,
        firstLineIndent=0.25 * inch,
        spaceAfter=8,
        splitLongWords=True,
    )
    story: list[object] = [
        Spacer(1, 2.15 * inch),
        Paragraph(escape(manuscript.title.strip()), title_style),
    ]
    if manuscript.author is not None:
        story.append(Paragraph(f"by {escape(manuscript.author.strip())}", byline_style))
    story.append(PageBreak())
    for index, scene in enumerate(manuscript.scenes):
        if index > 0:
            story.append(PageBreak())
        story.append(
            Paragraph(
                f"Scene {scene.scene_number}: {escape(scene.title.strip())}",
                scene_style,
            )
        )
        story.extend(
            Paragraph(escape(paragraph), body_style) for paragraph in _prose_paragraphs(scene.prose)
        )

    document = SimpleDocTemplate(
        output,
        pagesize=LETTER,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.85 * inch,
        bottomMargin=0.75 * inch,
        title=manuscript.title,
        author=manuscript.author or "Open Hollywood",
        subject="Open Hollywood prose manuscript",
        creator="Open Hollywood",
    )
    document.build(
        story,
        onFirstPage=_draw_pdf_footer,
        onLaterPages=_draw_pdf_footer,
        canvasmaker=_InvariantCanvas,
    )
    return output.getvalue()


def export_docx(manuscript: ProseManuscript) -> bytes:
    """Export a stable, editable US-Letter DOCX using the editorial preset."""
    document = create_document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    _configure_docx_styles(document)
    _configure_docx_metadata(document, manuscript)
    _add_docx_header_and_footer(document, manuscript.title)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(150)
    title.add_run(manuscript.title.strip())
    if manuscript.author is not None:
        byline = document.add_paragraph()
        byline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        byline.paragraph_format.space_before = Pt(8)
        run = byline.add_run(f"by {manuscript.author.strip()}")
        run.italic = True
    document.add_page_break()  # type: ignore[no-untyped-call]

    for index, scene in enumerate(manuscript.scenes):
        if index > 0:
            document.add_page_break()  # type: ignore[no-untyped-call]
        document.add_heading(f"Scene {scene.scene_number}: {scene.title.strip()}", level=1)
        for paragraph_text in _prose_paragraphs(scene.prose):
            document.add_paragraph(paragraph_text, style="Normal")

    raw = BytesIO()
    document.save(raw)
    return _canonicalize_docx(raw.getvalue())


def _configure_docx_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_tokens = (
        ("Title", 28, _DARK_BLUE, 0, 18),
        ("Heading 1", 16, _BLUE, 18, 10),
        ("Heading 2", 13, _BLUE, 12, 6),
        ("Heading 3", 12, _DARK_BLUE, 8, 4),
    )
    for name, size, color, before, after in heading_tokens:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Manuscript Header" not in styles:
        header_style = styles.add_style("Manuscript Header", WD_STYLE_TYPE.PARAGRAPH)
    else:
        header_style = styles["Manuscript Header"]
    header_style.font.name = "Calibri"
    header_style.font.size = Pt(8)
    header_style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _configure_docx_metadata(document: Document, manuscript: ProseManuscript) -> None:
    properties = document.core_properties
    properties.title = manuscript.title
    properties.author = manuscript.author or "Open Hollywood"
    properties.subject = "Open Hollywood prose manuscript"
    properties.keywords = "Open Hollywood, short prose, manuscript"
    properties.comments = "Deterministically rendered from immutable scene drafts."
    properties.created = _STATIC_TIMESTAMP
    properties.modified = _STATIC_TIMESTAMP
    properties.last_printed = _STATIC_TIMESTAMP
    properties.revision = 1


def _add_docx_header_and_footer(document: Document, title: str) -> None:
    section = document.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    header_paragraph.style = "Manuscript Header"
    header_paragraph.text = title.strip()

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.style = "Manuscript Header"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Open Hollywood  |  ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)  # noqa: SLF001


def _draw_pdf_footer(canvas: Canvas, document: SimpleDocTemplate) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColorRGB(0.4, 0.4, 0.4)
    canvas.drawCentredString(LETTER[0] / 2, 0.4 * inch, str(document.page))
    canvas.restoreState()


class _InvariantCanvas(Canvas):  # type: ignore[misc]
    """Force ReportLab's reproducible metadata and document identifier mode."""

    def __init__(self, *args: Any, **kwargs: Any) -> None:
        kwargs["invariant"] = 1
        super().__init__(*args, **kwargs)


def _prose_paragraphs(prose: str) -> tuple[str, ...]:
    normalized = prose.replace("\r\n", "\n").replace("\r", "\n").strip()
    return tuple(part.strip() for part in normalized.split("\n\n") if part.strip())


def _canonicalize_docx(payload: bytes) -> bytes:
    source = BytesIO(payload)
    output = BytesIO()
    with (
        ZipFile(source, "r") as archive,
        ZipFile(
            output,
            "w",
            compression=ZIP_DEFLATED,
            compresslevel=9,
        ) as canonical,
    ):
        for name in sorted(archive.namelist()):
            info = ZipInfo(filename=name, date_time=_ZIP_TIMESTAMP)
            info.compress_type = ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            canonical.writestr(info, archive.read(name))
    return output.getvalue()
